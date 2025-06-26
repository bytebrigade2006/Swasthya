# doc_generator.py - DEBUGGED VERSION
import google.generativeai as genai
from datetime import datetime
from docx import Document
import io
import json
import os
import logging

# Configure logging
logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

def init_gemini_client():
    """
    Initialize Gemini client with API key from environment variables
    Returns: GenerativeModel instance or None on failure
    """
    try:
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            logger.error("Gemini API key not found in environment variables")
            return None
        genai.configure(api_key=api_key)
        return genai.GenerativeModel('gemini-1.5-flash')
    except Exception as e:
        logger.error(f"Gemini initialization error: {str(e)}")
        return None

def clean_json_response(response_text):
    """
    Clean and extract JSON from Gemini response
    Args:
        response_text: Raw response string from Gemini
    Returns: JSON string or None if invalid
    """
    try:
        # Remove markdown code blocks
        response_text = response_text.strip()
        
        # FIX 1: Proper string handling for markdown removal
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        elif response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
        
        # Extract JSON content
        start = response_text.find('{')
        end = response_text.rfind('}') + 1
        if start == -1 or end == 0:
            logger.warning("No JSON found in response")
            return None
            
        return response_text[start:end]
    except Exception as e:
        logger.error(f"JSON cleaning error: {str(e)}")
        return None

def extract_information(document_text, document_type):
    """
    Extract structured information from document text using Gemini
    Args:
        document_text: Extracted text from document
        document_type: Type of healthcare document
    Returns: Dict of extracted data or None on failure
    """
    model = init_gemini_client()
    if not model:
        return None

    # FIX 2: Add input validation
    if not document_text or not document_text.strip():
        logger.error("Empty document text provided")
        return None

    prompt = f"""
    Extract relevant information from this {document_type} document:
    {document_text[:2000]}
    
    Return ONLY this JSON structure:
    {{
        "patient_name": "Full name or Not found",
        "policy_number": "Policy ID or Not found",
        "date_of_birth": "YYYY-MM-DD or Not found",
        "phone": "Phone number or Not found",
        "email": "Email or Not found",
        "address": "Full address or Not found",
        "diagnosis": "Medical diagnosis or Not found",
        "treatment": "Treatment details or Not found",
        "service_date": "YYYY-MM-DD or Not found",
        "provider_name": "Provider name or Not found",
        "claim_amount": "Numbers only or Not found",
        "insurance_company": "Insurance company or Not found"
    }}
    Return ONLY the JSON object.
    """

    try:
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.1,
                max_output_tokens=800,
                top_p=0.8,
                top_k=40
            )
        )
        
        if not response or not response.text:
            logger.error("Empty response from Gemini")
            return None
            
        json_str = clean_json_response(response.text)
        if not json_str:
            return None
            
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        logger.error(f"JSON parsing error: {str(e)}")
        logger.error(f"Raw response: {response.text if response else 'No response'}")
        return None
    except Exception as e:
        logger.error(f"Information extraction failed: {str(e)}")
        return None

def generate_document_content(document_type, patient_data, claim_details):
    """
    Generate document content using Gemini
    Args:
        document_type: Type of document to generate
        patient_data: Dict of patient information
        claim_details: Dict of claim details
    Returns: Generated text content or None on failure
    """
    model = init_gemini_client()
    if not model:
        return None

    # FIX 3: Add input validation and safe get operations
    if not patient_data or not claim_details:
        logger.error("Missing patient data or claim details")
        return None

    # FIX 4: Safe handling of amount formatting
    try:
        amount = float(claim_details.get('amount', 0))
    except (ValueError, TypeError):
        amount = 0.0

    prompt = f"""
    Generate a professional {document_type} with:
    Patient Information:
    - Name: {patient_data.get('name', 'Not provided')}
    - Policy: {patient_data.get('policy_number', 'Not provided')}
    - DOB: {patient_data.get('dob', 'Not provided')}
    - Contact: {patient_data.get('contact', 'Not provided')}
    
    Claim Details:
    - Service Date: {claim_details.get('service_date', 'Not provided')}
    - Diagnosis: {claim_details.get('diagnosis', 'Not provided')}
    - Treatment: {claim_details.get('treatment', 'Not provided')}
    - Amount: ₹{amount:,.2f}
    - Reason: {claim_details.get('reason', 'Not provided')}
    
    Include:
    1. Professional letter format
    2. Clear claim statement
    3. Medical justification
    4. Specific action request
    5. Professional closing
    
    Use Indian Rupees (₹) and Indian healthcare standards.
    """
    
    try:
        response = model.generate_content(
            prompt,
            generation_config=genai.types.GenerationConfig(
                temperature=0.3,
                max_output_tokens=1500,
                top_p=0.8,
                top_k=40
            )
        )
        
        # FIX 5: Better response validation
        if not response or not response.text:
            logger.error("Empty response from Gemini for content generation")
            return None
            
        return response.text.strip()
    except Exception as e:
        logger.error(f"Content generation failed: {str(e)}")
        return None

def create_word_document(content, doc_type, patient_name):
    """
    Create Word document from generated content
    Args:
        content: Generated document text
        doc_type: Type of document
        patient_name: Patient's name
    Returns: Bytes of Word document or None on failure
    """
    try:
        # FIX 6: Input validation
        if not content or not content.strip():
            logger.error("No content provided for document creation")
            return None
            
        doc = Document()
        
        # FIX 7: Proper header handling - headers need paragraphs added
        header_para = doc.sections[0].header.paragraphs[0]
        header_para.text = f"{doc_type} - {patient_name or 'Unknown Patient'}"
        
        # Title
        title = doc.add_heading(doc_type, 0)
        title.alignment = 1  # Center
        
        # Date
        date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = 2  # Right align
        
        # Content - FIX 8: Handle multi-line content properly
        content_lines = content.split('\n')
        for line in content_lines:
            if line.strip():  # Skip empty lines
                doc.add_paragraph(line.strip())
        
        # FIX 9: Proper footer handling
        footer_para = doc.sections[0].footer.paragraphs[0]
        footer_para.text = "Generated by AI Healthcare Document Generator"
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()
    except Exception as e:
        logger.error(f"Word document creation failed: {str(e)}")
        return None

# FIX 10: Add helper functions for better error handling and validation
def validate_patient_data(patient_data):
    """Validate patient data structure"""
    required_fields = ['name', 'policy_number', 'dob']
    if not isinstance(patient_data, dict):
        return False
    return any(patient_data.get(field) for field in required_fields)

def validate_claim_details(claim_details):
    """Validate claim details structure"""
    required_fields = ['service_date', 'diagnosis', 'amount']
    if not isinstance(claim_details, dict):
        return False
    return any(claim_details.get(field) for field in required_fields)

# FIX 11: Add main execution function with proper error handling
def main():
    """Main function for testing the document generator"""
    try:
        # Test data
        test_patient = {
            'name': 'John Doe',
            'policy_number': 'POL123456',
            'dob': '1990-01-15',
            'contact': 'john@example.com'
        }
        
        test_claim = {
            'service_date': '2024-01-20',
            'diagnosis': 'Hypertension',
            'treatment': 'Medication prescribed',
            'amount': 5000.00,
            'reason': 'Regular checkup and treatment'
        }
        
        if validate_patient_data(test_patient) and validate_claim_details(test_claim):
            content = generate_document_content("Insurance Claim Letter", test_patient, test_claim)
            if content:
                doc_bytes = create_word_document(content, "Insurance Claim Letter", test_patient['name'])
                if doc_bytes:
                    print("Document generated successfully!")
                    return doc_bytes
                else:
                    print("Failed to create Word document")
            else:
                print("Failed to generate content")
        else:
            print("Invalid test data")
    except Exception as e:
        logger.error(f"Main execution failed: {str(e)}")
        return None

if __name__ == "__main__":
    main()

def generate_document_content(document_type, patient_data, claim_details):
    """Generate specific content for each document type"""
    model = init_gemini_client()
    if not model:
        return None

    # Document-specific templates
    if document_type == "Insurance Claim Letter":
        prompt = f"""
            Write a complete, ready-to-send Insurance Claim Letter using the exact details below.
            Do NOT use any placeholders like [Your Name] or [Your Address].
            Fill in all sections with the provided data.

            Patient Name: {patient_data.get('patient_name', 'N/A')}
            Policy Number: {patient_data.get('policy_number', 'N/A')}
            Date of Birth: {patient_data.get('date_of_birth', 'N/A')}
            Phone: {patient_data.get('phone', 'N/A')}
            Email: {patient_data.get('email', 'N/A')}
            Address: {patient_data.get('address', 'N/A')}
            Service Date: {claim_details.get('service_date', 'N/A')}
            Diagnosis: {claim_details.get('diagnosis', 'N/A')}
            Treatment: {claim_details.get('treatment', 'N/A')}
            Claim Amount: ₹{claim_details.get('claim_amount', '0')}
            Provider Name: {claim_details.get('provider_name', 'N/A')}
            Insurance Company: {claim_details.get('insurance_company', 'N/A')}
            Reason: {claim_details.get('reason', 'N/A')}

            Write the letter as if you are the patient, with all the above details filled in the appropriate places. Do NOT use any placeholders or ask for missing information.
            """
        
    elif document_type == "Medical Report":
       prompt = f"""
            Write a comprehensive Medical Report using the exact details below.
            Do NOT use any placeholders like [Your Name] or [Your Address].
            Fill in all sections with the provided data.

            Patient Name: {patient_data.get('patient_name', 'N/A')}
            Date of Birth: {patient_data.get('date_of_birth', 'N/A')}
            Phone: {patient_data.get('phone', 'N/A')}
            Email: {patient_data.get('email', 'N/A')}
            Address: {patient_data.get('address', 'N/A')}
            Service Date: {claim_details.get('service_date', 'N/A')}
            Diagnosis: {claim_details.get('diagnosis', 'N/A')}
            Treatment: {claim_details.get('treatment', 'N/A')}
            Provider Name: {claim_details.get('provider_name', 'N/A')}
            Include medical history, examination findings, diagnosis, and treatment plan. Do NOT use any placeholders or ask for missing information.
            """
        
    elif document_type == "Appeal Letter":
        prompt = f"""
            Write a complete, ready-to-send Appeal Letter using the exact details below. 
            Do NOT use any placeholders like [Your Name] or [Your Address]. 
            Fill in all sections with the provided data.
            
            Patient Name: {patient_data.get('patient_name', 'N/A')}
            Policy Number: {patient_data.get('policy_number', 'N/A')}
            Original Claim Date: {claim_details.get('service_date', 'N/A')}
            Diagnosis: {claim_details.get('diagnosis', 'N/A')}
            Treatment: {claim_details.get('treatment', 'N/A')}
            Denied Amount: ₹{claim_details.get('claim_amount', '0')}
            Appeal Reason: {claim_details.get('reason', 'N/A')}
            Insurance Company: {claim_details.get('insurance_company', 'N/A')}
            Phone: {patient_data.get('phone', 'N/A')}
            Email: {patient_data.get('email', 'N/A')}
            Provider Name: {claim_details.get('provider_name', 'N/A')}
            Date of Birth: {patient_data.get('date_of_birth', 'N/A')}
            Address: {patient_data.get('address', 'N/A')}

            Write the letter as if you are the patient, with all the above details filled in the appropriate places. Do NOT use any placeholders or ask for missing information.
            """

    
    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        logger.error(f"Content generation failed: {str(e)}")
        return None
